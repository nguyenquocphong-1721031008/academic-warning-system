from docx import Document


def add_api(doc: Document, method: str, path: str, desc: str, access: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{method} {path}").bold = True
    doc.add_paragraph(f"Mo ta: {desc}")
    doc.add_paragraph(f"Phan quyen: {access}")


def main() -> None:
    doc = Document()
    doc.add_heading("Mo ta chi tiet cac API hien co", level=1)
    doc.add_paragraph(
        "Tai lieu tong hop cac endpoint backend, chuc nang va doi tuong duoc phep truy cap."
    )

    doc.add_heading("1. Authentication", level=2)
    add_api(doc, "POST", "/api/auth/login", "Dang nhap, cap access token va refresh token.", "Public")
    add_api(doc, "POST", "/api/auth/refresh", "Lam moi access token bang refresh token.", "User da dang nhap")
    add_api(doc, "POST", "/api/auth/logout", "Dang xuat, huy refresh token hien tai.", "User da dang nhap")
    add_api(doc, "GET", "/api/auth/me", "Lay thong tin tai khoan hien tai.", "User da dang nhap")

    doc.add_heading("2. Warnings", level=2)
    add_api(doc, "GET", "/api/warnings", "Danh sach canh bao hoc vu (co loc, phan trang).", "Admin/Faculty manager")
    add_api(doc, "GET", "/api/warnings/analytics", "Thong ke canh bao theo hoc ky, khoa, lop.", "Admin/Faculty manager")
    add_api(doc, "GET", "/api/warnings/export", "Xuat danh sach canh bao ra CSV.", "Admin/Faculty manager")
    add_api(doc, "GET", "/api/warnings/analysis/{student_code}", "Phan tich tong hop rule-based + ML cho 1 MSSV.", "Public")
    add_api(doc, "GET", "/api/warnings/{student_code}", "Tra cuu trang thai canh bao theo MSSV.", "Public/Internal")

    doc.add_heading("3. Machine Learning", level=2)
    add_api(doc, "GET", "/api/ml/status", "Kiem tra model da load, threshold, metrics.", "Internal")
    add_api(doc, "POST", "/api/ml/predict", "Du bao rui ro canh bao hoc vu cho MSSV.", "Internal")
    add_api(doc, "POST", "/api/ml/train", "Train 1 model canh bao hoc vu.", "Admin")
    add_api(doc, "POST", "/api/ml/train-all", "Train tat ca model trong 1 lan goi API.", "Admin")

    doc.add_heading("4. Admin", level=2)
    add_api(doc, "POST", "/api/admin/users", "Tao tai khoan nguoi dung.", "Admin")
    add_api(doc, "GET", "/api/admin/users", "Lay danh sach nguoi dung.", "Admin")
    add_api(doc, "DELETE", "/api/admin/users/{user_id}", "Xoa nguoi dung.", "Admin")
    add_api(doc, "POST", "/api/admin/users/{user_id}/reset-password", "Reset mat khau nguoi dung.", "Admin")
    add_api(doc, "POST", "/api/admin/faculties", "Tao khoa.", "Admin")
    add_api(doc, "GET", "/api/admin/faculties", "Lay danh sach khoa.", "Admin")
    add_api(doc, "PUT", "/api/admin/faculties/{faculty_id}", "Cap nhat khoa.", "Admin")
    add_api(doc, "DELETE", "/api/admin/faculties/{faculty_id}", "Xoa khoa.", "Admin")
    add_api(doc, "POST", "/api/admin/warning-rule-sets", "Tao bo luat canh bao.", "Admin")
    add_api(doc, "GET", "/api/admin/warning-rule-sets", "Lay danh sach bo luat.", "Admin")
    add_api(doc, "PUT", "/api/admin/warning-rule-sets/{rule_set_id}", "Cap nhat bo luat.", "Admin")
    add_api(doc, "POST", "/api/admin/warning-rule-sets/{rule_set_id}/toggle", "Bat/tat bo luat.", "Admin")
    add_api(doc, "POST", "/api/admin/warning-rules", "Tao rule canh bao.", "Admin")
    add_api(doc, "GET", "/api/admin/warning-rules", "Lay danh sach rule.", "Admin")
    add_api(doc, "GET", "/api/admin/warning-rule-sets/{rule_set_id}/rules", "Lay rule theo bo luat.", "Admin")
    add_api(doc, "PUT", "/api/admin/warning-rules/{rule_id}", "Cap nhat rule.", "Admin")
    add_api(doc, "DELETE", "/api/admin/warning-rules/{rule_id}", "Xoa rule.", "Admin")
    add_api(doc, "POST", "/api/admin/warnings/regenerate", "Tai tao du lieu canh bao hoc vu.", "Admin")
    add_api(doc, "POST", "/api/admin/send-warning-email", "Gui email canh bao tu dong theo MSSV.", "Admin")

    doc.add_heading("5. Faculty Manager", level=2)
    add_api(doc, "GET", "/api/faculty-manager/students", "Lay danh sach sinh vien trong khoa.", "Faculty manager")
    add_api(doc, "GET", "/api/faculty-manager/warnings", "Lay danh sach canh bao trong khoa.", "Faculty manager")

    doc.add_heading("6. Scores", level=2)
    add_api(doc, "POST", "/api/scores/import", "Import bang diem tu file Excel.", "Admin")

    output = "d:\\DoAn\\academic-warning-backend\\API_ChucNang_ChiTiet.docx"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
